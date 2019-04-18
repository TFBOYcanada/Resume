
# coding: utf-8

# In[2]:


from docx import Document
from nltk import word_tokenize, pos_tag
from nltk.stem import WordNetLemmatizer
import nltk
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
import re


# In[3]:


MONTHS = {
        1: ["jan", "january", "01"],
        2: ["feb", "february", "02"],
        3: ["mar", "march", "03"],
        4: ["apr", "april", "04"],
        5: ["may", "05"],
        6: ["jun", "june", "06"],
        7: ["jul", "july", "07"],
        8: ["aug", "august", "08"],
        9: ["sep", "sept", "september", "09"],
        10: ["oct", "october", "10"],
        11: ["nov", "november", "11"],
        12: ["dec", "december", "12"]
    }
EXAMS = {"gre", "gmat"}
CERTIFICATES = {
    'cpa', 'ca', 'cicpa', 'hkicpa', 'acca', 'cga', 'cma', 'cta', 'cisa', 'cia', 'cima', 'frm', 'cfa', 'csc', 'cbv',
    'soa', 'cda', 'caa', 'rfc', 'cwm', 'cqf', 'cva', 'fia', 'fiaa', 'fsa', 'phr', 'ihri', 'cir', 'chrp', 'hrpa'
               }
LEADERSHIP = {'president', 'board', 'vp', 'lead', 'found'}


# In[ ]:


document = Document('OSFI.docx')


# # General Check
# ## Font Size, Tense, Format, etc.

# In[ ]:


def check_font_size(document):
    abnormal_size_char_count = 0
    total_char_count = 0
    for t in document.paragraphs:
        if t.text != '' and len(t.runs)>0:
            size = t.style.font.size if t.style.font.size is not None else t.runs[0].font.size
            text_len = len(t.text)
            if size > 152400 or size < 127000:
                abnormal_size_char_count += text_len
            total_char_count += text_len
    if (abnormal_size_char_count/total_char_count) > 0.1:
        print("""按照您简历的具体内容丰富程度，字体大小为10-12号。小于10号的字体，容易让HR or Hiring manager在一份简历<30秒阅读的规则下，看不清楚；如果字体大于12号，可能容易变成2页简历""")
        return True
    else:
        return False


# In[ ]:


def check_tense(document):
    count_past = 0
    count_future = 0
    count_present = 0
    past_list = ["VBD", "VBN"]
    present_list = ["VBP", "VBZ", "VBG"]
    future_list = ["VBC", "VBF"]
    
    for t in document.paragraphs:
        for _, tag in pos_tag(word_tokenize(t.text)): 
            if tag in past_list:
                count_past += 1
            if tag in present_list: 
                count_present += 1
            if tag in future_list:
                count_future += 1
    if count_future > 0 or count_present * 2 > count_past:
        print("""过去的经历应该使用过去时态， 当前的工作可以用现在时态""")
        return True
    return False


# # Helper Function break sentences into common parts in resume

# In[ ]:


def break_parts(document):
    parts_dic = {
        "education": [],
        "experience": [],
        "activity": [],
        "skill": [],
        "certification": [],
        "summary": [],
        "awards": [],
        "projects": [],
        "research": [], 
        "language": []
    }
    
    lemmatizer = WordNetLemmatizer()
    key_words_dic = {
        "education": [lemmatizer.lemmatize("education")],
        "experience": [lemmatizer.lemmatize("experience")],
        "activity": [lemmatizer.lemmatize("activity"), lemmatizer.lemmatize("extracurricular"), lemmatizer.lemmatize("leadership")],
        "skill": [lemmatizer.lemmatize("skill")],
        "certification": [lemmatizer.lemmatize("certification"), lemmatizer.lemmatize("certificate")],
        "summary": [lemmatizer.lemmatize("summary")],
        "awards": [lemmatizer.lemmatize("awards"), lemmatizer.lemmatize("honors")],
        "projects": [lemmatizer.lemmatize("projects")],
        "research": [lemmatizer.lemmatize("research")],
        "language": [lemmatizer.lemmatize("language")]
    }
    
    buffer = []
    key = None
    tokens_set = set()
    for t in document.paragraphs:
        tokens = []
        
        for token in word_tokenize(t.text):
#             if bool(re.match("^[A-Za-z0-9.]*$", token)):
            tokens.append(lemmatizer.lemmatize(token.lower()))
            tokens_set.add(lemmatizer.lemmatize(token.lower()))
                
#         print(tokens)
        
        for sec in key_words_dic.keys():
            for key_word in key_words_dic[sec]:
                if key_word in tokens and len(tokens) <= 4:
                    if key is not None:
                        parts_dic[key].extend(buffer)
                    buffer = []
                    key = sec
        if tokens != []:
            buffer.append(tokens)
    if key is not None:
        parts_dic[key].extend(buffer)
    return parts_dic, tokens_set
                


# In[ ]:


parts, token_set = break_parts(document)


# In[ ]:


def check_parts(parts):
    if len(parts['education']) <= 0:
        print("""Please add education!""")
        return True
    if len(parts['experience']) <= 0:
        print("""Please add experience!""")
        return True
    if len(parts['skill']) <= 0:
        print("""Please add skills!""")
        return True
    return False


# # Check Part 1: Summary

# In[ ]:


def check_summary(part_dic):
    if len(part_dic["summary"]) == 0:
        return False
    count_summary = 0
    count_total = 0
    for key in part_dic.keys():
        if key == "summary":
            for token in part_dic[key]:
                count_summary += len(token)
        for token in part_dic[key]:
            count_total += len(token)
    if (count_summary / count_total) > 0.1:
        print("""通常建议，个人自我评价不应该占用整份简历的10%的篇幅，因为您有什么潜力，什么特长优势，是结合您的个人经历体现的。""")
        return True
    return False


# # Check Part 2: Education

# In[ ]:


edu = parts["education"]


# ## helper function to get graduation year

# In[ ]:


def get_graduation_year(edu_list):
    graduation_year = 0
    for lst in edu_list:
        for token in lst:
            if token.isdigit():
                if int(token) > 1950:
                    graduation_year = max(graduation_year, int(token))
            elif token == "present":
                graduation_year = 2019
    return graduation_year


# In[ ]:


grad_year = get_graduation_year(edu)


# ## helper function to check if month in a line with year

# In[ ]:


def check_date_format(line):
    month = None
    year = None
    for token in line:
        if token.isdigit() and int(token) > 1950:
            year = int(token)
        for key in MONTHS.keys():
            if token in MONTHS[key]:
                month = key
    if month is not None and year is not None:
        return False, (month, year)
    elif month is None and year is None:
        return False, (None, None)
    else:
        return True, (month, year)


# In[ ]:


def check_education_date_format(edu_list):
    one_valid = False
    for line in edu_list:
        check_format, content = check_date_format(line)
        if check_format:
            print("""通常建议把（预计）毕业年月份写出来，告知招聘官你可以全职入职的时间，这个其实是非常重要的一个信息
            写法1:　University of Toronto at scarborough (expected graduated on May 2020)
            写法2:  University of Western Ontario (graudated on May 2018)""")
            return True
        if None not in content:
            one_valid = True
    if not one_valid:
        print("""通常建议把（预计）毕业年月份写出来，告知招聘官你可以全职入职的时间，这个其实是非常重要的一个信息
            写法1:　University of Toronto at scarborough (expected graduated on May 2020)
            写法2:  University of Western Ontario (graudated on May 2018)""")
        return True
    else:
        return False


# In[ ]:


def check_gpa(edu_list, grad_year=2019):
    if grad_year > 2017:
        gpa = False
        for line in edu_list:
            if "gpa" in line:
                gpa = True

            for token in line:
                if re.match("^\d+?\.\d+?$", token) is not None: #or (len(token.split("/") > 1) and re.match("^\d+?\.\d+?$", token.split("/")[0]) is not None):
                    if float(token) < 3.0: 
                        print("""通常填写两位数，并且把总分写出来（也可以四舍五入写一位数），成绩优异可以描述名次。尚未毕业则写当前已知的GPA 如果GPA较低，可以计算一个"专业GPA"，如果GPA低于3.0可以不写 

                                写法1：GPA：3.76 / 4.00
                                写法2：GPA：3.86 / 4.00 (Top 5%) 
                                写法3：GPA：3.0 (Core GPA: 3.5) 
                                文科/文理学院的 GPA 在 3.6 以上属于比较“高”； 理科/工程学院的 GPA 在 3.2 以上属于比较“高”。 如果你的 GPA 看起来比较低但实际上在整个专业中排名很靠前，可以写一个括号注明，比如：GPA:3.0 (全专业前 10%)""")
                        return True
                    elif float(token) <= 4.0 and float(token) >= 3.0:
                        gpa = True
        if not gpa:
            print("""通常填写两位数，并且把总分写出来（也可以四舍五入写一位数），成绩优异可以描述名次。尚未毕业则写当前已知的GPA 如果GPA较低，可以计算一个"专业GPA"，如果GPA低于3.0可以不写 

                    写法1：GPA：3.76 / 4.00
                    写法2：GPA：3.86 / 4.00 (Top 5%) 
                    写法3：GPA：3.0 (Core GPA: 3.5) 
                    文科/文理学院的 GPA 在 3.6 以上属于比较“高”； 理科/工程学院的 GPA 在 3.2 以上属于比较“高”。 如果你的 GPA 看起来比较低但实际上在整个专业中排名很靠前，可以写一个括号注明，比如：GPA:3.0 (全专业前 10%)""")
            return True
    return False


# In[ ]:


def check_relevant_course(edu_list, grad_year=2019):
    if grad_year == 2019:
        course = False
        for line in edu_list:
            for token in line: 
                if "course" in token:
                    course = True
        if not course:
            print("""应届求职或找实习：如果经历不多，内容不够一页纸，适当添加相关课程，建议不要超过1行，无需写每课成绩。 已工作跳槽：建议不写课程。如有1段以上教育经历，只写最高学历的、跟岗位相关的课程即可。 考研或学术：以重要性和成绩罗列跟申请专业/科研项目相关的课程，可写成绩，挑选填写不用把所有课都列出。 
                写法1（求职）：Relevant Coursework: Accounting & Finance, Leadership in Business, Financial Economics 
                写法2（学术）：Relevant Coursework: Data Analytics (95), Intro to Statistics (93), Big Data (92)""")
            return True
    return False
    


# In[ ]:


def check_exam_score(edu_list):
    exam = False
    for line in edu_list:
        for token in line: 
            if token in EXAMS:
                exam = True
    if not exam:
        print("""对于MBA或者想加入MBB or 2nd Tier战略咨询的申请者，建议写GMAT/GRE成绩
            写法1 (美国求职者): GMAT>=720 (xx% percentile) 有竞争力（也有特例低于720分的成功入选者）
            写法2 (加拿大求职者): GMAT>=700 (xx% percentile) 有竞争力（也有特例低于700分的成功入选者）
            写法3 (美国求职者): GRE>330 (xx% percentile) 有竞争力（也有特例低于330分的成功入选者）
            写法4 (加拿大求职者): GRE>300 (xx% percentile) 有竞争力（也有特例低于300分的成功入选者）

            对于GMAT/GRE低于benchmark,　请咨询Career Check咨询行业顾问""")
        return True
    return False


# In[ ]:


def check_education_length(edu_list, token_set, grad_year=2019):
    edu_sum = 0
    for line in edu_list:
        edu_sum += len(line)
    if grad_year >= 2018 and edu_sum / len(token_set) > 0.3:
        print("""too much education, please add more experience!""")
        return True
    if grad_year < 2018 and edu_sum / len(token_set) > 0.2:
        print("""too much education, please add more experience!""")
        return True
    return False


# # Check Part 3: Awards and Honors

# In[ ]:


awards = parts["awards"]


# In[ ]:


def check_awards(awards_list, grad_year=2019):
    if grad_year == 2019:
        if len(awards_list) <= 0:
            print("awards missing!")
            return True
    return False
            


# ## helper function to check entries in chronological order

# In[ ]:


def check_chronological(part_list):
    last_year = 2019
    last_month = 12
    for line in part_list:
        _, content = check_date_format(line)
        if content[1] is not None and content[1] > last_year:
            print(last_year, last_month, content)
            print("chronological order 1!")
            return True
        elif content[1] is not None and content[1] == last_year:
            if content[0] is not None and content[0] > last_month:
                print(last_year, last_month, content)
                print("chronological order 2!")
                return True
        if content[1] is not None:
            last_year = content[1]
        if content[0] is not None:
            last_month = content[0]
    return False


# # Check Part 4: Experience

# In[ ]:


def check_quantify_result(experience_list):
    total_experience = 0
    quantify = 0
    result = False
    for line in experience_list:
        check_format, content = check_date_format(line)
        if not check_format and None not in content:
            total_experience += 1
            quantify += int(result)
            result = False
        else:
            if not result:
                for token in line:
                    if token.isdigit():
                        result = True
    if total_experience > 0 and quantify / total_experience < 0.5:
        print("total experience {}, quantified results {}. add more quantifiable results!".format(total_experience, quantify))
        return True
    else:
        return False


# In[ ]:


def check_experience_date_format(experience_list):
    one_valid = False
    for line in experience_list:
        check_format, content = check_date_format(line)
        if check_format:
            print(line)
            print("""时间是简历 中必不可少的要素，教育经历、工作经历、荣誉奖励和技能证书都需标注时间，日期的书写格式有很多种，我们参考了MS Office中的日期格式，其中适合简历书写表达得主要有以下几种：

                    1-2012年6月
                    2-2012/06
                    3-2012.06
                    4-June 12
                    5. Jun. 2012
                    6-06/2012
                    无论你选择哪种格式，关键是时间格式要统一

                    常见的错误格式有以下几种情况，您可以参考如下：
                    错误一：两段时间的间隔线不统一，存在一长一短的现象
                    错误二：未精确到月
                    错误三：两段时间间隔为“~”或者“/”，与其他时间段的表达方式不统一
                    错误四：其他时间表述如2003/3，月份前面未加0
                    错误五：时间没有右对齐
                    错误六：20004，笔误""")
            return True
        if None not in content:
            one_valid = True
    if not one_valid:
        print("""时间是简历 中必不可少的要素，教育经历、工作经历、荣誉奖励和技能证书都需标注时间，日期的书写格式有很多种，我们参考了MS Office中的日期格式，其中适合简历书写表达得主要有以下几种：

                1-2012年6月
                2-2012/06
                3-2012.06
                4-June 12
                5. Jun. 2012
                6-06/2012
                无论你选择哪种格式，关键是时间格式要统一

                常见的错误格式有以下几种情况，您可以参考如下：
                错误一：两段时间的间隔线不统一，存在一长一短的现象
                错误二：未精确到月
                错误三：两段时间间隔为“~”或者“/”，与其他时间段的表达方式不统一
                错误四：其他时间表述如2003/3，月份前面未加0
                错误五：时间没有右对齐
                错误六：20004，笔误""")
        return True
    else:
        return False


# # Check Part 5: Certification

# In[ ]:


def check_certificates(part_dic, tokens_set):
    if len(part_dic["certification"]) == 0:
        for shorthands in CERTIFICATES:
            if shorthands in tokens_set:
                return False
        print("""Add certifications if applicable!""")
        return True
    return False


# # Check Part 6: Activities and Extracurricular

# In[ ]:


def check_activities(acticity_lst):
    if len(acticity_lst) != 0:
        total_experience = 0
        impact = 0
        result = False
        for line in acticity_lst:
            check_format, content = check_date_format(line)
            if not check_format and None not in content:
                total_experience += 1
                impact += int(result)
                result = False
            else:
                if not result:
                    for token in line:
                        if token.isdigit() or token in LEADERSHIP:
                            result = True
        if total_experience > 0 and impact / total_experience < 0.8:
            print("total experience {}, visible impact {}. add more quantifiable results!".format(total_experience, impact))
            return True
       
    return False


# # Check Part 7: Skills

# In[ ]:


def check_skills(skill_list):
    if len(skill_list) <= 0:
        print("""Please add skill if applicable!""")
        return True
    return False


# In[ ]:


def check_languages(language_list):
    if len(language_list) <= 0:
        print("""Please add language if applicable!""")
        return True
    return False


# In[ ]:


check_font_size(document)
check_tense(document)
check_summary(parts)
check_parts(parts)
check_education_date_format(parts['education'])
check_gpa(parts['education'])
check_relevant_course(parts['education'], grad_year)
check_exam_score(parts['education'])
check_education_length(parts['education'], token_set, grad_year)
check_awards(parts['awards'], grad_year)
check_experience_date_format(parts['experience'])
check_quantify_result(parts['experience'])
check_certificates(parts, token_set)
check_activities(parts['activity'])
check_skills(parts['skill'])
check_languages(parts['language'])

